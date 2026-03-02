"""
Word Export API - Génération de documents Word à partir de tables HTML
Utilise python-docx pour une génération fiable côté serveur
Supporte les couleurs de risque (Probabilité, Impact, Criticité)
"""

import logging
import re
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional
from io import BytesIO
from datetime import datetime

# Import python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with: pip install python-docx")

logger = logging.getLogger("word-export")

router = APIRouter(prefix="/api/word", tags=["Word Export"])

# Patterns pour détecter les colonnes de risque
RISK_PATTERNS = {
    'probabilite': re.compile(r'probabilit[eé]|prob|obab', re.IGNORECASE),
    'impact': re.compile(r'impact|mpa|impa', re.IGNORECASE),
    'criticite': re.compile(r'criticit[eé]|ritici', re.IGNORECASE)
}

# Couleurs de risque (RGB)
RISK_COLORS = {
    'Faible': {'bg': '28A745', 'rgb': RGBColor(40, 167, 69), 'text_rgb': RGBColor(255, 255, 255)},    # Vert
    'Moyen': {'bg': 'FFC107', 'rgb': RGBColor(255, 193, 7), 'text_rgb': RGBColor(0, 0, 0)},          # Jaune
    'Eleve': {'bg': 'DC3545', 'rgb': RGBColor(220, 53, 69), 'text_rgb': RGBColor(255, 255, 255)}     # Rouge
}


class TableData(BaseModel):
    """Données d'une table"""
    headers: List[str]
    rows: List[List[str]]


class ExportWordRequest(BaseModel):
    """Requête d'export Word"""
    tables: List[TableData]
    filename: Optional[str] = None


def set_cell_shading(cell, color: str):
    """Appliquer une couleur de fond à une cellule"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def normalize_risk_value(value: str) -> Optional[str]:
    """
    Normalise une valeur de risque (F/M/E ou Faible/Moyen/Eleve)
    Retourne le niveau normalisé ou None
    """
    if not value:
        return None
    
    v = str(value).strip().lower()
    
    # Correspondances exactes
    if v in ['f', 'faible', 'low', 'bas']:
        return 'Faible'
    if v in ['m', 'moyen', 'medium', 'mod', 'modéré']:
        return 'Moyen'
    if v in ['e', 'é', 'eleve', 'élevé', 'elevé', 'high', 'haut', 'h']:
        return 'Eleve'
    
    # Recherche partielle
    if 'faible' in v or 'low' in v or 'bas' in v:
        return 'Faible'
    if 'moyen' in v or 'medium' in v or 'modéré' in v:
        return 'Moyen'
    if 'eleve' in v or 'élevé' in v or 'high' in v or 'haut' in v:
        return 'Eleve'
    
    # Première lettre
    first_char = v[0] if v else ''
    if first_char == 'f':
        return 'Faible'
    if first_char == 'm':
        return 'Moyen'
    if first_char in ['e', 'é', 'h']:
        return 'Eleve'
    
    return None


def find_risk_columns(headers: List[str]) -> dict:
    """Trouve les index des colonnes de risque"""
    indexes = {'probabilite': -1, 'impact': -1, 'criticite': -1}
    
    for idx, header in enumerate(headers):
        h = str(header).lower().strip()
        if RISK_PATTERNS['probabilite'].search(h):
            indexes['probabilite'] = idx
        if RISK_PATTERNS['impact'].search(h):
            indexes['impact'] = idx
        if RISK_PATTERNS['criticite'].search(h):
            indexes['criticite'] = idx
    
    return indexes


def create_word_document(tables_data: List[TableData]) -> BytesIO:
    """
    Créer un document Word avec les tables fournies
    Style: en-têtes rouge bordeaux (#800020), texte blanc, bordures noires
    Couleurs de risque: Faible=Vert, Moyen=Jaune, Élevé=Rouge
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="python-docx non disponible. Installez avec: pip install python-docx"
        )
    
    doc = Document()
    
    # Couleurs
    burgundy = "800020"
    
    for table_index, table_data in enumerate(tables_data):
        # Ajouter un espace entre les tables (sauf avant la première)
        if table_index > 0:
            doc.add_paragraph()
        
        # Calculer le nombre de colonnes
        num_cols = len(table_data.headers) if table_data.headers else 0
        if num_cols == 0 and table_data.rows:
            num_cols = max(len(row) for row in table_data.rows)
        
        if num_cols == 0:
            continue
        
        # Nombre total de lignes (en-tête + données)
        num_rows = 1 + len(table_data.rows) if table_data.headers else len(table_data.rows)
        
        if num_rows == 0:
            continue
        
        # Détecter les colonnes de risque
        risk_columns = find_risk_columns(table_data.headers) if table_data.headers else {}
        risk_col_indexes = set()
        for key, idx in risk_columns.items():
            if idx >= 0:
                risk_col_indexes.add(idx)
        
        logger.info(f"📊 Table {table_index + 1}: {num_cols} colonnes, {num_rows} lignes, colonnes risque: {risk_columns}")
        
        # Créer la table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        row_index = 0
        
        # Ajouter les en-têtes si présents
        if table_data.headers:
            header_row = table.rows[0]
            for col_index, header_text in enumerate(table_data.headers):
                if col_index < num_cols:
                    cell = header_row.cells[col_index]
                    
                    # Vider et créer un nouveau run pour le style
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(str(header_text) if header_text else "")
                    
                    # Style de l'en-tête
                    set_cell_shading(cell, burgundy)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            
            row_index = 1
        
        # Ajouter les données
        for data_row in table_data.rows:
            if row_index < num_rows:
                table_row = table.rows[row_index]
                for col_index, cell_text in enumerate(data_row):
                    if col_index < num_cols:
                        cell = table_row.cells[col_index]
                        
                        # Vérifier si c'est une colonne de risque
                        is_risk_col = col_index in risk_col_indexes
                        risk_level = None
                        
                        if is_risk_col and cell_text:
                            risk_level = normalize_risk_value(cell_text)
                        
                        # Vider le contenu existant et créer un nouveau paragraphe
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(str(cell_text) if cell_text else "")
                        
                        # Appliquer le style
                        if risk_level and risk_level in RISK_COLORS:
                            # Style de risque avec couleur
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_cell_shading(cell, RISK_COLORS[risk_level]['bg'])
                            run.font.bold = True
                            run.font.color.rgb = RISK_COLORS[risk_level]['text_rgb']
                            run.font.size = Pt(9)
                        else:
                            # Style normal
                            run.font.size = Pt(9)
                
                row_index += 1
        
        # Ajuster les largeurs de colonnes
        for col_index in range(num_cols):
            # Calculer la largeur basée sur le contenu
            max_width = len(table_data.headers[col_index]) if table_data.headers and col_index < len(table_data.headers) else 5
            for row in table_data.rows:
                if col_index < len(row):
                    max_width = max(max_width, len(row[col_index]))
            
            # Limiter la largeur
            width_cm = min(max(max_width * 0.2, 2), 8)
            
            for row in table.rows:
                if col_index < len(row.cells):
                    row.cells[col_index].width = Cm(width_cm)
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@router.post("/export")
async def export_word(request: ExportWordRequest):
    """
    Exporter des tables en document Word
    
    Reçoit une liste de tables avec leurs en-têtes et données,
    génère un document Word avec le style approprié.
    """
    try:
        logger.info(f"📄 Export Word: {len(request.tables)} tables")
        
        # Créer le document
        doc_buffer = create_word_document(request.tables)
        
        # Générer le nom de fichier
        if request.filename:
            filename = request.filename
        else:
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"claraverse_export_{timestamp}.docx"
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export Word: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/health")
async def health_check():
    """Vérifier si le service Word est disponible"""
    return {
        "status": "ok" if DOCX_AVAILABLE else "unavailable",
        "docx_available": DOCX_AVAILABLE,
        "message": "python-docx disponible" if DOCX_AVAILABLE else "Installez python-docx: pip install python-docx"
    }


# ============================================================================
# EXPORT RAPPORT D'AUDIT STRUCTURÉ
# ============================================================================

# Colonnes possibles par étape de mission d'audit
AUDIT_COLUMNS = {
    'synthese_frap': ['no', 'intitule', 'observation', 'constat', 'risque', 'recommandation'],
    'rapport_provisoire': ['no', 'intitule', 'observation', 'constat', 'risque', 'recommandation', 'commentaire_audite'],
    'rapport_final': ['no', 'intitule', 'observation', 'constat', 'risque', 'recommandation', 
                      'commentaire_audite', 'commentaire_auditeur', 'plan_action', 'responsable', 'delai']
}

# Mapping des noms de colonnes (insensible à la casse et aux accents)
COLUMN_MAPPING = {
    'no': ['no', 'n°', 'numero', 'numéro', 'num', '#'],
    'intitule': ['intitule', 'intitulé', 'titre', 'libelle', 'libellé', 'designation', 'désignation'],
    'observation': ['observation', 'observations', 'analyse', 'contexte'],
    'constat': ['constat', 'constats', 'finding', 'findings', 'anomalie', 'anomalies'],
    'risque': ['risque', 'risques', 'risk', 'risks', 'incidence', 'incidences', 'impact', 'impacts'],
    'recommandation': ['recommandation', 'recommandations', 'recommendation', 'recommendations', 'preconisation', 'préconisation'],
    'commentaire_audite': ['commentaire_audite', 'commentaire audite', 'commentaire audité', 'reponse_audite', 'réponse audité', 'reponse audité'],
    'commentaire_auditeur': ['commentaire_auditeur', 'commentaire auditeur', 'reponse_auditeur', 'réponse auditeur'],
    'plan_action': ['plan_action', 'plan action', 'plan d\'action', 'actions', 'action'],
    'responsable': ['responsable', 'responsables', 'owner', 'pilote'],
    'delai': ['delai', 'délai', 'echeance', 'échéance', 'deadline', 'date']
}


class AuditReportRow(BaseModel):
    """Données d'une ligne de rapport d'audit"""
    no: Optional[str] = None
    intitule: Optional[str] = None
    observation: Optional[str] = None
    constat: Optional[str] = None
    risque: Optional[str] = None
    recommandation: Optional[str] = None
    commentaire_audite: Optional[str] = None
    commentaire_auditeur: Optional[str] = None
    plan_action: Optional[str] = None
    responsable: Optional[str] = None
    delai: Optional[str] = None


class AuditReportRequest(BaseModel):
    """Requête d'export rapport d'audit"""
    headers: List[str]
    rows: List[List[str]]
    etape: Optional[str] = None  # synthese_frap, rapport_provisoire, rapport_final
    reference: Optional[str] = None
    filename: Optional[str] = None


def normalize_column_name(header: str) -> Optional[str]:
    """
    Normalise un nom de colonne pour le mapper aux colonnes standard
    """
    if not header:
        return None
    
    h = header.lower().strip()
    # Supprimer les accents et caractères spéciaux
    h = h.replace('é', 'e').replace('è', 'e').replace('ê', 'e')
    h = h.replace('à', 'a').replace('â', 'a')
    h = h.replace('ô', 'o').replace('î', 'i').replace('û', 'u')
    h = h.replace('_', ' ').replace('-', ' ')
    
    for standard_name, variants in COLUMN_MAPPING.items():
        for variant in variants:
            variant_normalized = variant.lower().replace('é', 'e').replace('è', 'e').replace('ê', 'e')
            variant_normalized = variant_normalized.replace('à', 'a').replace('â', 'a')
            variant_normalized = variant_normalized.replace('ô', 'o').replace('î', 'i').replace('û', 'u')
            if h == variant_normalized or variant_normalized in h or h in variant_normalized:
                return standard_name
    
    return None


def detect_audit_stage(headers: List[str]) -> str:
    """
    Détecte l'étape de mission d'audit basée sur les colonnes présentes
    """
    normalized_headers = [normalize_column_name(h) for h in headers]
    normalized_set = set(h for h in normalized_headers if h)
    
    # Vérifier si c'est un rapport final (a toutes les colonnes)
    if 'plan_action' in normalized_set or 'responsable' in normalized_set or 'delai' in normalized_set:
        return 'rapport_final'
    
    # Vérifier si c'est un rapport provisoire (a commentaire_audite)
    if 'commentaire_audite' in normalized_set:
        return 'rapport_provisoire'
    
    # Par défaut, c'est une synthèse FRAP
    return 'synthese_frap'


def create_audit_report_document(request: AuditReportRequest) -> BytesIO:
    """
    Créer un document Word structuré pour un rapport d'audit
    Format: sections numérotées avec Observation, Constat, Risque, Recommandation, etc.
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="python-docx non disponible. Installez avec: pip install python-docx"
        )
    
    doc = Document()
    
    # Détecter l'étape de mission
    etape = request.etape or detect_audit_stage(request.headers)
    logger.info(f"📋 Export Rapport Audit - Étape détectée: {etape}")
    
    # Mapper les colonnes
    column_indexes = {}
    for idx, header in enumerate(request.headers):
        normalized = normalize_column_name(header)
        if normalized:
            column_indexes[normalized] = idx
    
    logger.info(f"📊 Colonnes mappées: {column_indexes}")
    
    # Couleurs - Bleu professionnel
    blue_rgb = RGBColor(30, 58, 138)  # #1E3A8A - Bleu foncé
    
    # Titre du rapport
    title = doc.add_heading('RAPPORT D\'AUDIT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = blue_rgb
    
    # Sous-titre avec l'étape
    etape_labels = {
        'synthese_frap': 'Synthèse des FRAP',
        'rapport_provisoire': 'Rapport Provisoire',
        'rapport_final': 'Rapport Final'
    }
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(etape_labels.get(etape, 'Rapport'))
    run.font.size = Pt(14)
    run.font.italic = True
    
    # Référence si fournie
    if request.reference:
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_para.add_run(f"Référence: {request.reference}")
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espace
    
    # Traiter chaque ligne comme une constatation
    for row_idx, row in enumerate(request.rows):
        # Numéro de la constatation
        no = row[column_indexes['no']] if 'no' in column_indexes and column_indexes['no'] < len(row) else str(row_idx + 1)
        
        # Intitulé
        intitule = row[column_indexes['intitule']] if 'intitule' in column_indexes and column_indexes['intitule'] < len(row) else f"Constatation {no}"
        
        # Titre de la section
        section_title = doc.add_heading(f"{no}. {intitule}", level=1)
        for run in section_title.runs:
            run.font.color.rgb = blue_rgb
        
        # Observation/Analyse
        if 'observation' in column_indexes and column_indexes['observation'] < len(row):
            observation = row[column_indexes['observation']]
            if observation and observation.strip():
                obs_title = doc.add_paragraph()
                run = obs_title.add_run("Observation/Analyse")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                obs_content = doc.add_paragraph(observation)
                obs_content.paragraph_format.left_indent = Cm(0.5)
        
        # Constat
        if 'constat' in column_indexes and column_indexes['constat'] < len(row):
            constat = row[column_indexes['constat']]
            if constat and constat.strip():
                const_title = doc.add_paragraph()
                run = const_title.add_run("Constat")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                const_content = doc.add_paragraph(constat)
                const_content.paragraph_format.left_indent = Cm(0.5)
        
        # Risque(s) / Incidence(s)
        if 'risque' in column_indexes and column_indexes['risque'] < len(row):
            risque = row[column_indexes['risque']]
            if risque and risque.strip():
                risk_title = doc.add_paragraph()
                run = risk_title.add_run("Risque(s) / Incidence(s)")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                # Traiter les risques comme liste si séparés par \n ou -
                risques = risque.replace('\\n', '\n').split('\n')
                for r in risques:
                    r = r.strip()
                    if r:
                        if r.startswith('-'):
                            r = r[1:].strip()
                        risk_item = doc.add_paragraph(f"• {r}")
                        risk_item.paragraph_format.left_indent = Cm(0.5)
        
        # Recommandation
        if 'recommandation' in column_indexes and column_indexes['recommandation'] < len(row):
            recommandation = row[column_indexes['recommandation']]
            if recommandation and recommandation.strip():
                rec_title = doc.add_paragraph()
                run = rec_title.add_run("Recommandation")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                # Traiter les recommandations comme liste si séparées par \n ou -
                recs = recommandation.replace('\\n', '\n').split('\n')
                for r in recs:
                    r = r.strip()
                    if r:
                        if r.startswith('-'):
                            r = r[1:].strip()
                        rec_item = doc.add_paragraph(f"• {r}")
                        rec_item.paragraph_format.left_indent = Cm(0.5)
        
        # Commentaire audité (rapport provisoire et final)
        if 'commentaire_audite' in column_indexes and column_indexes['commentaire_audite'] < len(row):
            commentaire = row[column_indexes['commentaire_audite']]
            if commentaire and commentaire.strip():
                com_title = doc.add_paragraph()
                run = com_title.add_run("Commentaire de l'audité")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                com_content = doc.add_paragraph(commentaire)
                com_content.paragraph_format.left_indent = Cm(0.5)
        
        # Commentaire auditeur (rapport final)
        if 'commentaire_auditeur' in column_indexes and column_indexes['commentaire_auditeur'] < len(row):
            commentaire = row[column_indexes['commentaire_auditeur']]
            if commentaire and commentaire.strip():
                com_title = doc.add_paragraph()
                run = com_title.add_run("Commentaire de l'auditeur")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                com_content = doc.add_paragraph(commentaire)
                com_content.paragraph_format.left_indent = Cm(0.5)
        
        # Plan d'action (rapport final)
        if 'plan_action' in column_indexes and column_indexes['plan_action'] < len(row):
            plan = row[column_indexes['plan_action']]
            if plan and plan.strip():
                plan_title = doc.add_paragraph()
                run = plan_title.add_run("Plan d'action")
                run.font.bold = True
                run.font.color.rgb = blue_rgb
                run.font.size = Pt(11)
                
                plan_content = doc.add_paragraph(plan)
                plan_content.paragraph_format.left_indent = Cm(0.5)
        
        # Tableau récapitulatif Actions/Responsables/Délais (rapport final)
        has_action_table = False
        responsable = None
        delai = None
        
        if 'responsable' in column_indexes and column_indexes['responsable'] < len(row):
            responsable = row[column_indexes['responsable']]
            if responsable and responsable.strip():
                has_action_table = True
        
        if 'delai' in column_indexes and column_indexes['delai'] < len(row):
            delai = row[column_indexes['delai']]
            if delai and delai.strip():
                has_action_table = True
        
        if has_action_table:
            # Créer un petit tableau pour les actions
            action_table = doc.add_table(rows=2, cols=3)
            action_table.style = 'Table Grid'
            
            # En-têtes
            headers_row = action_table.rows[0]
            for idx, header_text in enumerate(['Actions', 'Responsables', 'Délai']):
                cell = headers_row.cells[idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(header_text)
                run.font.bold = True
                run.font.size = Pt(9)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, "1E3A8A")  # Bleu foncé
                run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Données
            data_row = action_table.rows[1]
            plan_text = row[column_indexes['plan_action']] if 'plan_action' in column_indexes and column_indexes['plan_action'] < len(row) else ""
            
            for idx, value in enumerate([plan_text or "", responsable or "", delai or ""]):
                cell = data_row.cells[idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(value))
                run.font.size = Pt(9)
        
        # Séparateur entre les constatations
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run("─" * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        doc.add_paragraph()
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@router.post("/export-rapport")
async def export_audit_report(request: AuditReportRequest):
    """
    Exporter une table d'audit en rapport Word structuré
    
    Génère un document Word avec les sections:
    - Observation/Analyse
    - Constat
    - Risque(s) / Incidence(s)
    - Recommandation
    - Commentaire audité (si présent)
    - Commentaire auditeur (si présent)
    - Plan d'action avec tableau Actions/Responsables/Délai (si présent)
    """
    try:
        logger.info(f"📋 Export Rapport Audit: {len(request.rows)} constatations")
        logger.info(f"📊 Colonnes: {request.headers}")
        
        # Créer le document
        doc_buffer = create_audit_report_document(request)
        
        # Générer le nom de fichier
        if request.filename:
            filename = request.filename
        else:
            etape = request.etape or detect_audit_stage(request.headers)
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"rapport_audit_{etape}_{timestamp}.docx"
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export rapport audit: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# EXPORT FRAP INDIVIDUELLE (Multi-tables à une colonne)
# ============================================================================

class FrapTable(BaseModel):
    """Une table FRAP individuelle (une colonne)"""
    header: str
    content: str


class FrapExportRequest(BaseModel):
    """Requête d'export FRAP individuelle"""
    tables: List[FrapTable]
    reference: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    filename: Optional[str] = None


def create_frap_document(request: FrapExportRequest) -> BytesIO:
    """
    Créer un document Word pour une FRAP individuelle
    Structure: plusieurs tables à une colonne (Intitulé, Observation, Constat, Risque, Recommandation)
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="python-docx non disponible. Installez avec: pip install python-docx"
        )
    
    doc = Document()
    
    # Couleur bleue professionnelle
    blue_rgb = RGBColor(30, 58, 138)  # #1E3A8A
    blue_hex = "1E3A8A"
    
    # Extraire les métadonnées et le contenu
    metadata = {}
    content_sections = {}
    
    for table in request.tables:
        header_lower = table.header.lower().strip()
        
        # Métadonnées
        if 'etape' in header_lower or 'mission' in header_lower:
            metadata['etape'] = table.content
        elif 'norme' in header_lower:
            metadata['norme'] = table.content
        elif 'methode' in header_lower or 'méthode' in header_lower:
            metadata['methode'] = table.content
        elif 'reference' in header_lower or 'référence' in header_lower:
            metadata['reference'] = table.content
        # Contenu principal
        elif 'intitul' in header_lower or 'titre' in header_lower:
            content_sections['intitule'] = table.content
        elif 'observation' in header_lower or 'analyse' in header_lower:
            content_sections['observation'] = table.content
        elif 'constat' in header_lower:
            content_sections['constat'] = table.content
        elif 'risque' in header_lower or 'incidence' in header_lower:
            content_sections['risque'] = table.content
        elif 'recomm' in header_lower or 'préconisation' in header_lower:
            content_sections['recommandation'] = table.content
    
    # Utiliser les paramètres de la requête si fournis
    if request.reference:
        metadata['reference'] = request.reference
    if request.norme:
        metadata['norme'] = request.norme
    if request.methode:
        metadata['methode'] = request.methode
    
    logger.info(f"📋 Export FRAP - Métadonnées: {metadata}")
    logger.info(f"📊 Sections de contenu: {list(content_sections.keys())}")
    
    # Titre du document
    title = doc.add_heading('FICHE DE RÉVÉLATION ET D\'ANALYSE DE PROBLÈME', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(16)
    
    # Sous-titre (FRAP)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("FRAP")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = blue_rgb
    
    doc.add_paragraph()
    
    # Tableau des métadonnées
    if metadata:
        meta_table = doc.add_table(rows=len(metadata), cols=2)
        meta_table.style = 'Table Grid'
        
        meta_items = [
            ('Référence', metadata.get('reference', '')),
            ('Étape de mission', metadata.get('etape', 'FRAP')),
            ('Norme', metadata.get('norme', '')),
            ('Méthode', metadata.get('methode', ''))
        ]
        
        row_idx = 0
        for label, value in meta_items:
            if value:
                row = meta_table.rows[row_idx] if row_idx < len(meta_table.rows) else meta_table.add_row()
                
                # Cellule label
                cell_label = row.cells[0]
                cell_label.text = ""
                para = cell_label.paragraphs[0]
                run = para.add_run(label)
                run.font.bold = True
                run.font.size = Pt(10)
                set_cell_shading(cell_label, blue_hex)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Cellule valeur
                cell_value = row.cells[1]
                cell_value.text = ""
                para = cell_value.paragraphs[0]
                run = para.add_run(str(value))
                run.font.size = Pt(10)
                
                row_idx += 1
        
        # Supprimer les lignes vides
        while len(meta_table.rows) > row_idx:
            meta_table._tbl.remove(meta_table.rows[-1]._tr)
    
    doc.add_paragraph()
    
    # Intitulé (titre de la constatation)
    if 'intitule' in content_sections:
        intitule_title = doc.add_paragraph()
        run = intitule_title.add_run("INTITULÉ")
        run.font.bold = True
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(12)
        
        intitule_content = doc.add_paragraph()
        run = intitule_content.add_run(content_sections['intitule'])
        run.font.size = Pt(11)
        run.font.bold = True
        intitule_content.paragraph_format.left_indent = Cm(0.5)
        
        doc.add_paragraph()
    
    # Observation/Analyse
    if 'observation' in content_sections:
        obs_title = doc.add_paragraph()
        run = obs_title.add_run("OBSERVATION / ANALYSE")
        run.font.bold = True
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(12)
        
        # Traiter le contenu avec les <br>
        obs_text = content_sections['observation'].replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        obs_content = doc.add_paragraph(obs_text)
        obs_content.paragraph_format.left_indent = Cm(0.5)
        
        doc.add_paragraph()
    
    # Constat
    if 'constat' in content_sections:
        const_title = doc.add_paragraph()
        run = const_title.add_run("CONSTAT")
        run.font.bold = True
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(12)
        
        const_text = content_sections['constat'].replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        const_content = doc.add_paragraph(const_text)
        const_content.paragraph_format.left_indent = Cm(0.5)
        
        doc.add_paragraph()
    
    # Risque(s) / Incidence(s)
    if 'risque' in content_sections:
        risk_title = doc.add_paragraph()
        run = risk_title.add_run("RISQUE(S) / INCIDENCE(S)")
        run.font.bold = True
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(12)
        
        # Traiter les risques comme liste
        risk_text = content_sections['risque'].replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        risques = risk_text.split('\n')
        for r in risques:
            r = r.strip()
            if r:
                if r.startswith('-'):
                    r = r[1:].strip()
                risk_item = doc.add_paragraph(f"• {r}")
                risk_item.paragraph_format.left_indent = Cm(0.5)
        
        doc.add_paragraph()
    
    # Recommandation
    if 'recommandation' in content_sections:
        rec_title = doc.add_paragraph()
        run = rec_title.add_run("RECOMMANDATION")
        run.font.bold = True
        run.font.color.rgb = blue_rgb
        run.font.size = Pt(12)
        
        # Traiter les recommandations comme liste
        rec_text = content_sections['recommandation'].replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        recs = rec_text.split('\n')
        for r in recs:
            r = r.strip()
            if r:
                if r.startswith('-'):
                    r = r[1:].strip()
                rec_item = doc.add_paragraph(f"• {r}")
                rec_item.paragraph_format.left_indent = Cm(0.5)
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@router.post("/export-frap")
async def export_frap(request: FrapExportRequest):
    """
    Exporter une FRAP individuelle en document Word structuré
    
    Structure attendue: plusieurs tables à une colonne chacune
    - Table métadonnées (Étape, Norme, Méthode, Référence)
    - Table Intitulé
    - Table Observation
    - Table Constat
    - Table Risque
    - Table Recommandation
    """
    try:
        logger.info(f"📋 Export FRAP: {len(request.tables)} tables")
        
        # Créer le document
        doc_buffer = create_frap_document(request)
        
        # Générer le nom de fichier
        if request.filename:
            filename = request.filename
        else:
            ref = request.reference or "FRAP"
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"frap_{ref}_{timestamp}.docx"
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export FRAP: {e}")
        raise HTTPException(status_code=500, detail=str(e))
